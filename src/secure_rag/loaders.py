from __future__ import annotations

import csv
import json
from collections.abc import Iterable, Iterator
from pathlib import Path
from typing import Any

from secure_rag.schema import Metadata, SourceDocument

SUPPORTED_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".pdf",
    ".docx",
    ".json",
    ".yaml",
    ".yml",
    ".csv",
    ".xlsx",
}


def load_path(path: str | Path) -> list[SourceDocument]:
    resolved = Path(path)
    if resolved.is_dir():
        return list(load_directory(resolved))
    if not resolved.exists():
        raise FileNotFoundError(resolved)

    suffix = resolved.suffix.lower()
    if suffix in {".txt", ".md", ".markdown"}:
        return [_load_text(resolved)]
    if suffix == ".pdf":
        return _load_pdf(resolved)
    if suffix == ".docx":
        return [_load_docx(resolved)]
    if suffix == ".json":
        return _load_json(resolved)
    if suffix in {".yaml", ".yml"}:
        return _load_yaml(resolved)
    if suffix == ".csv":
        return _load_csv(resolved)
    if suffix == ".xlsx":
        return _load_xlsx(resolved)
    raise ValueError(f"Unsupported file extension: {suffix}")


def load_directory(path: str | Path) -> Iterator[SourceDocument]:
    root = Path(path)
    for child in sorted(root.rglob("*")):
        if child.is_file() and child.suffix.lower() in SUPPORTED_EXTENSIONS:
            yield from load_path(child)


def _base_metadata(path: Path, **extra: str | int | float | bool | None) -> Metadata:
    return {"file_name": path.name, "file_type": path.suffix.lower().lstrip("."), **extra}


def _load_text(path: Path) -> SourceDocument:
    return SourceDocument(
        path=path,
        text=path.read_text(encoding="utf-8"),
        metadata=_base_metadata(path),
    )


def _load_pdf(path: Path) -> list[SourceDocument]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("Install pypdf to load PDF files") from exc

    reader = PdfReader(str(path))
    documents: list[SourceDocument] = []
    for page_index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            documents.append(
                SourceDocument(
                    path=path,
                    text=text,
                    metadata=_base_metadata(path, page=page_index),
                )
            )
    return documents


def _load_docx(path: Path) -> SourceDocument:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError("Install python-docx to load Word files") from exc

    document = docx.Document(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    return SourceDocument(path=path, text="\n".join(paragraphs), metadata=_base_metadata(path))


def _load_json(path: Path) -> list[SourceDocument]:
    data = json.loads(path.read_text(encoding="utf-8"))
    return _structured_documents(path, data)


def _load_yaml(path: Path) -> list[SourceDocument]:
    try:
        import yaml
    except ImportError as exc:
        raise RuntimeError("Install pyyaml to load YAML files") from exc

    data = yaml.safe_load(path.read_text(encoding="utf-8"))
    return _structured_documents(path, data)


def _load_csv(path: Path) -> list[SourceDocument]:
    rows: list[SourceDocument] = []
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        for index, row in enumerate(csv.DictReader(handle), start=1):
            text = _record_to_text(row)
            rows.append(
                SourceDocument(
                    path=path,
                    text=text,
                    metadata=_base_metadata(path, record_index=index),
                )
            )
    return rows


def _load_xlsx(path: Path) -> list[SourceDocument]:
    try:
        import openpyxl
    except ImportError as exc:
        raise RuntimeError("Install openpyxl to load Excel files") from exc

    workbook = openpyxl.load_workbook(path, data_only=True, read_only=True)
    documents: list[SourceDocument] = []
    for sheet in workbook.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue
        headers = [
            str(value) if value is not None else f"column_{i}"
            for i, value in enumerate(rows[0])
        ]
        for row_index, row in enumerate(rows[1:], start=2):
            record = {
                headers[column_index]: value
                for column_index, value in enumerate(row)
                if value not in (None, "")
            }
            if record:
                documents.append(
                    SourceDocument(
                        path=path,
                        text=_record_to_text(record),
                        metadata=_base_metadata(path, worksheet=sheet.title, row=row_index),
                    )
                )
    return documents


def _structured_documents(path: Path, data: Any) -> list[SourceDocument]:
    if isinstance(data, list):
        return [
            SourceDocument(
                path=path,
                text=_record_to_text(item),
                metadata=_base_metadata(path, record_index=index),
            )
            for index, item in enumerate(data, start=1)
        ]
    return [SourceDocument(path=path, text=_record_to_text(data), metadata=_base_metadata(path))]


def _record_to_text(record: Any) -> str:
    if isinstance(record, dict):
        return "\n".join(f"{key}: {_scalar_to_text(value)}" for key, value in _flatten(record))
    if isinstance(record, list):
        return "\n".join(_scalar_to_text(value) for value in record)
    return _scalar_to_text(record)


def _flatten(value: dict[str, Any], prefix: str = "") -> Iterable[tuple[str, Any]]:
    for key, item in value.items():
        name = f"{prefix}.{key}" if prefix else str(key)
        if isinstance(item, dict):
            yield from _flatten(item, name)
        else:
            yield name, item


def _scalar_to_text(value: Any) -> str:
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=True, sort_keys=True)
    return "" if value is None else str(value)
